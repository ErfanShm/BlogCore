import os
import re
import json
import logging
import markdown as md
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaFileUpload
from googleapiclient.http import MediaIoBaseUpload

import requests
from openai import APIConnectionError

# Load env
load_dotenv()
logger = logging.getLogger(__name__)

# Google Drive configs
GOOGLE_CREDENTIALS_FILE = os.getenv("GOOGLE_CREDENTIALS_FILE", "credentials.json")
GOOGLE_DRIVE_FOLDER_ID  = os.getenv("GOOGLE_DRIVE_FOLDER_ID")
API_KEY_GOOGLE          = os.getenv("API_KEY_GOOGLE")
ENGINE_ID_GOOGLE        = os.getenv("ENGINE_ID_GOOGLE")

# KWRank configs
KWRANK_CLIENT_ID        = os.getenv("KWRANK_CLIENT_ID")
KWRANK_SECRET_ID        = os.getenv("KWRANK_SECRET_ID")
KWRANK_AUTH_TOKEN       = None # Will be populated after successful authorization

from .llm_clients import get_openai_clients

# ── Helpers ────────────────────────────────────────────────────────────────

def get_kwrank_auth_token():
    global KWRANK_AUTH_TOKEN
    if KWRANK_AUTH_TOKEN:
        return KWRANK_AUTH_TOKEN

    if not KWRANK_CLIENT_ID or not KWRANK_SECRET_ID:
        logger.error("KWRANK_CLIENT_ID or KWRANK_SECRET_ID not set.")
        return None

    auth_url = "https://kwrank.ir/api/v1/authorize"
    headers = {'Content-Type': 'application/x-www-form-urlencoded'}
    data = {
        'client_id': KWRANK_CLIENT_ID,
        'secret_id': KWRANK_SECRET_ID
    }

    try:
        response = requests.post(auth_url, headers=headers, data=data)
        response.raise_for_status() # Raise an exception for HTTP errors
        result = response.json()
        if result.get("result") and result.get("access_token"):
            KWRANK_AUTH_TOKEN = result["access_token"]
            logger.info("Successfully obtained KWRank access token.")
            return KWRANK_AUTH_TOKEN
        else:
            logger.error("Failed to obtain KWRank access token: %s", result.get("message", "Unknown error."))
            return None
    except requests.exceptions.RequestException as e:
        logger.error("Error connecting to KWRank authorization API: %s", e)
        return None

def count_words_fa(text: str) -> int:
    return len(re.findall(r"\w+", text)) if text else 0

def upload_as_google_doc(local_filepath, article_title, folder_id):
    """
    Upload a .docx file to Google Drive as a Google Doc, then delete the local file.
    Uses MediaIoBaseUpload so the file handle is closed before os.remove().
    """
    try:
        # 1) Authenticate
        creds = service_account.Credentials.from_service_account_file(
            GOOGLE_CREDENTIALS_FILE,
            scopes=["https://www.googleapis.com/auth/drive"]
        )
        drive = build("drive", "v3", credentials=creds)

        # 2) Prepare metadata
        metadata = {
            "name":       article_title,
            "parents":    [folder_id],
            "mimeType":   "application/vnd.google-apps.document"
        }

        # 3) Open the file in a with–block
        with open(local_filepath, "rb") as fh:
            media = MediaIoBaseUpload(
                fh,
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                resumable=True
            )
            # 4) Upload
            file = drive.files().create(
                body=metadata,
                media_body=media,
                fields="id"
            ).execute()

        # 5) Now the 'fh' is closed—safe to delete
        # os.remove(local_filepath)
        return file.get("id")

    except Exception as e:
        logger.error(f"❌ Error uploading to Google Drive: {e}")
        return None

def create_word_document(markdown_text: str, file_title: str) -> str:
    safe = re.sub(r'[\\/*?:"<>|]', "", file_title)
    fname = f"مقاله - {safe}.docx"
    outdir = "data/generated_docs"
    os.makedirs(outdir, exist_ok=True)
    fullpath = os.path.join(outdir, fname)

    doc = docx.Document()
    style = doc.styles["Normal"]
    style.font.name = "Vazirmatn"; style.font.size = Pt(12)
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.sections[0].right_to_left = True

    in_table = False
    table = None

    for line in markdown_text.splitlines():
        line = line.rstrip()
        if not line:
            continue

        if line.startswith("|") and line.endswith("|"):
            cells = [c.strip() for c in line.split("|")[1:-1]]
            if not in_table:
                in_table = True
                table = doc.add_table(rows=1, cols=len(cells))
                table.style = "Table Grid"
                for i, txt in enumerate(cells):
                    table.rows[0].cells[i].text = txt
            else:
                row = table.add_row().cells
                for i, txt in enumerate(cells):
                    row[i].text = txt
            continue
        else:
            in_table = False

        if line.startswith("### "):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith("* "):
            p = doc.add_paragraph(line[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(re.sub(r"^\d+\.\s", "", line), style="List Number")
        elif line.startswith("> "):
            p = doc.add_paragraph(line[2:]).italic = True
        else:
            p = doc.add_paragraph()
            for part in re.split(r"(\*\*.*?\*\*)", line):
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2]); run.bold = True
                else:
                    p.add_run(part)
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.save(fullpath)
    return fullpath

# ── LLM Wrapper ────────────────────────────────────────────────────────────

def get_kwrank_keyword_suggestions(keyword: str):
    token = get_kwrank_auth_token()
    if not token:
        return {"error": "Could not get KWRank authorization token.", "status": 500}

    api_url = "https://kwrank.ir/api/v1/keyword-suggester"
    headers = {
        'Authorization': f'Bearer {token}',
        'Content-Type': 'application/x-www-form-urlencoded'
    }
    data = {
        'keyword': keyword
    }

    try:
        response = requests.post(api_url, headers=headers, data=data)
        response.raise_for_status()
        result = response.json()

        if result.get("result"):
            return {"suggested_keywords": result.get("suggested_keywords", []), "status": 200}
        else:
            return {"error": result.get("message", "Unknown KWRank API error."), "status": result.get("status", 500)}
    except requests.exceptions.RequestException as e:
        logger.error("Error calling KWRank keyword suggester API: %s", e)
        return {"error": f"Error calling KWRank keyword suggester API: {e}", "status": 500}

def generate_content_bundle_with_llm(prompt_text: str, temperature: float, expect_json: bool = False):
    """
    Calls OpenAI /chat/completions with function-calling ("auto").
    Falls back to regex if no function_call is returned.
    """
    client, FUNCTIONS = get_openai_clients()
    if client is None:
        msg = "OpenAI client not initialized."
        logger.error(msg)
        return {"error": msg}

    messages = [
        {"role": "system", "content": "You are a helpful assistant."},
        {"role": "user",   "content": prompt_text}
    ]

    try:
        # ask the model to auto‐decide between content vs. function call
        resp = client.chat.completions.create(
            model="gemini-2.5-flash-preview-05-20",
            messages=messages,
            temperature=temperature,
            functions=FUNCTIONS,
            function_call="auto"       # ← let the model choose
        )

        choice = resp.choices[0].message

        # 1) if function_call happened, parse arguments
        if choice.function_call and choice.function_call.arguments:
            return json.loads(choice.function_call.arguments)

        # 2) else, log raw and try regex fallback
        raw = choice.content or ""
        logger.error("LLM did not call function—raw content:\n%s", raw[:500])

        # try fenced JSON first
        m = re.search(r"```json\s*([\s\S]+?)\s*```", raw, flags=re.IGNORECASE)
        if m:
            js = m.group(1)
        else:
            # fallback to first {…} block
            fb = re.search(r"(\{(?:[^{}]|(?R))*\})", raw, flags=re.DOTALL)
            if fb:
                js = fb.group(1)
            else:
                return {"error": "No JSON found in assistant response."}

        try:
            return json.loads(js)
        except json.JSONDecodeError as je:
            return {"error": f"Regex‐extracted JSON parse error: {je}"}

    except APIConnectionError as ace:
        logger.error("OpenAI connection failed: %s", ace)
        return {"error": "Could not connect to OpenAI API."}
    except Exception as e:
        logger.exception("Unexpected error in LLM call:")
        return {"error": str(e)}

# ── Google Search (unchanged) ───────────────────────────────────────────────

def perform_Google_Search(query: str):
    try:
        url    = "https://www.googleapis.com/customsearch/v1"
        params = {"key": API_KEY_GOOGLE, "cx": ENGINE_ID_GOOGLE, "q": query, "num": 5}
        r      = requests.get(url, params=params)
        r.raise_for_status()
        return r.json().get("items", [])
    except Exception as e:
        logger.error("Google Search error: %s", e)
        return []